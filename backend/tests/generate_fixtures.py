import os
import json
import pandas as pd
import docx

def main() -> None:
    """Generate mock sample files for testing document parsing."""
    fixtures_dir = os.path.join(os.path.dirname(__file__), "fixtures")
    os.makedirs(fixtures_dir, exist_ok=True)

    # 1. Plain Text
    txt_path = os.path.join(fixtures_dir, "sample.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("Hello World from a plain text file.\nThis is the second line.")

    # 2. Markdown
    md_path = os.path.join(fixtures_dir, "sample.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# Sample Title\n\nThis is a paragraph in markdown.\n\n- Item 1\n- Item 2")

    # 3. JSON
    json_path = os.path.join(fixtures_dir, "sample.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({"title": "Sample JSON", "items": ["a", "b", "c"], "nested": {"key": "value"}}, f, indent=2)

    # 4. HTML
    html_path = os.path.join(fixtures_dir, "sample.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write("<!DOCTYPE html><html><head><title>Sample HTML</title></head><body><h1>Hello HTML</h1><p>This is a paragraph of HTML text.</p></body></html>")

    # 5. CSV
    csv_path = os.path.join(fixtures_dir, "sample.csv")
    df = pd.DataFrame({
        "Name": ["Alice", "Bob", "Charlie"],
        "Age": [25, 30, 35],
        "City": ["New York", "London", "Paris"]
    })
    df.to_csv(csv_path, index=False)

    # 6. XLSX
    xlsx_path = os.path.join(fixtures_dir, "sample.xlsx")
    df.to_excel(xlsx_path, index=False)

    # 7. DOCX
    docx_path = os.path.join(fixtures_dir, "sample.docx")
    doc = docx.Document()
    doc.add_heading("Sample DOCX Title", 0)
    doc.add_paragraph("This is a paragraph in the docx document.")
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Header A"
    hdr_cells[1].text = "Header B"
    row_cells = table.rows[1].cells
    row_cells[0].text = "Data A1"
    row_cells[1].text = "Data B1"
    doc.save(docx_path)

    # 8. PDF (Write a minimal valid PDF file)
    pdf_path = os.path.join(fixtures_dir, "sample.pdf")
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> >>\nendobj\n"
        b"4 0 obj\n<< /Length 50 >>\n"
        b"stream\n"
        b"BT\n/F1 12 Tf\n72 712 Td\n(Hello PDF World Page 1) Tj\nET\n"
        b"endstream\nendobj\n"
        b"xref\n0 5\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000056 00000 n \n"
        b"0000000111 00000 n \n"
        b"0000000257 00000 n \n"
        b"trailer\n<< /Size 5 /Root 1 0 R >>\n"
        b"startxref\n356\n%%EOF\n"
    )
    with open(pdf_path, "wb") as f:
        f.write(pdf_content)

    # 9. PNG Image
    png_path = os.path.join(fixtures_dir, "sample.png")
    try:
        from PIL import Image, ImageDraw
        img = Image.new("RGB", (400, 300), color="white")
        draw = ImageDraw.Draw(img)
        # Draw text
        draw.text((20, 20), "Sales Report 2026", fill="black")
        draw.text((20, 50), "Revenue up by 15 percent", fill="black")
        # Draw a bar chart
        draw.rectangle([50, 100, 100, 250], fill="red", outline="black")
        draw.rectangle([120, 80, 170, 250], fill="green", outline="black")
        img.save(png_path)
    except Exception as e:
        print(f"Failed to generate PNG fixture: {e}")

    print("Successfully generated all mock sample fixtures.")


if __name__ == "__main__":
    main()
